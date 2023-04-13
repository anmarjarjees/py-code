"""
Link: https://python-docx.readthedocs.io/en/latest/
"""
# Creating and writing on MS Word Document
from docx import Document

# This class "Inches" is to be used with images/pictures to change their sizes
from docx.shared import Inches


# Creating a new document object
document = Document()

# Creating a the new document file on the disk
# document.save("demo.docx")

# Adding contents to this document
"""
We can add: heading, paragraph, page break, table, picture
Link: https://python-docx.readthedocs.io/en/latest/user/quickstart.html
"""
# Adding Title and Headings:
document.add_heading("Python Course",0) # Title 
document.add_heading("Python with MS Word") # default heading1
document.add_heading("Using Docx library",2) # heading2# 
# Or:
document.add_heading('Using Docx library', level=2)

"""
We cannot modify the file when it's open:
PermissionError: [Errno 13] Permission denied: 'demo.docx'
"""
# document.save("demo.docx")

# Adding paragraphs:
p= document.add_paragraph("When there is a will, there is a way.")

# Keep adding the previous paragraphs with styles:
p.add_run(" Yes, it's true!").bold = True
p.add_run(" Just keep going and keep moving on.").italic=True
p.add_run(" That's it.")
"""
When there is a will, there is a way. It's true! Just keep going and keep moving on.
"""

# document.save("demo.docx")

# let's create a function named "save" for short :-)
def save():
    document.save("demo.docx")

"""
Paragraph styles in default template:
-------------------------------------
Link: https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html?highlight=Intense%20Quote#paragraph-styles-in-default-template
"""

"""
.add_paragraph("Any text you put here", style="?")
"""
document.add_paragraph('Here is my intense quote!', style='Intense Quote')

document.add_paragraph(
    'unordered list item1', style='List Bullet'
)

document.add_paragraph(
    'unordered list item2', style='List Bullet'
)

document.add_paragraph(
    'unordered list item3', style='List Bullet'
)

# save()


document.add_paragraph(
    'ordered list item1', style='List Number'
)

document.add_paragraph(
    'ordered list item2', style='List Number'
)

document.add_paragraph(
    'ordered list item3', style='List Number'
)

# save()

"""
Adding images
Link: https://python-docx.readthedocs.io/en/latest/user/quickstart.html?highlight=add_paragraph#adding-a-picture
"""
# By default, the added image appears at native size. 
# document.add_picture('py-programming.jpg')
"""
NOTE: the image will be very big!
Link: https://python-docx.readthedocs.io/en/latest/user/quickstart.html?highlight=add_paragraph#image-size
"""

# save()


# To get the image the size you want, 
# you can specify either its width or height in convenient units, 
# like inches or centimeters:
# we need to import "Inches"
document.add_picture('py-programming.jpg', width=Inches(5.0))

# save()

"""
Adding a table:
Link: https://python-docx.readthedocs.io/en/latest/user/quickstart.html#adding-a-table
"""

document.add_heading("Table1:",2)

table1 = document.add_table(rows=3, cols=4)
# Tables have several properties and methods

# you can always access a cell by its row and column indicies:
cell1 = table1.cell(0,0) # the right-hand cell in the top row of the table 

# Once you have a cell, you can put something in it:
cell1.text = 'Java, PHP, and Python'

cell2 = table1.cell(1,0)
cell2.text = 'JavaScript, ASP.NET'

# save()

"""
- The .rows property of a table provides access to individual rows, 
each of which has a .cells property. 
- The .cells property on both Row and Column supports indexed access, 
like a list:
"""
row1 = table1.rows[2]

row1.cells[0].text = 'Python'
row1.cells[1].text = 'PHP'
row1.cells[2].text = 'Java'

# save()

document.add_paragraph("")
document.add_heading("Table2:",2)

table2 = document.add_table(rows=1, cols=3)

# Adding cells for the first row (header row cells):

# the code in two lines
first_row = table2.rows[0]
heading_cells = first_row.cells

# Or just in one line:
hdr_cells = table2.rows[0].cells

# Adding contents in each cell of the first row:
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

"""
- The .rows and .columns collections on a table are iterable, 
- you can use them directly in a for loop. 
- Same with the .cells sequences on a row or column:
"""

# Nested loop:
for row in table1.rows: # loop through all the table row(s)
    for cell in row.cells: # loop through all the cells in each row
        print(cell.text)

save()

# adding page break
document.add_page_break()

document.add_heading("Table3:",2)
table3 = document.add_table(rows=2, cols=3)

# creating a simple list (array)
t_headings = ["First Name","Last Name","Program"]

# the main list rows => all the table rows
rows = [
    # sub-lists => the list of the content of each row
    ['Alex','Chow','CSTN'],
    ['Alex','White','CMPG'],
    ['John','Doe','CSTN'],
    ['Alice','Brown','CMPG'],
    ['Bob','Johnson','CSTN']
]

# 3 lines of code!
"""
table3.rows[0].cells[0].text = t_headings[0] # First Name
table3.rows[0].cells[1].text = t_headings[1] # Last Name
table3.rows[0].cells[2].text = t_headings[2] # Program
"""


# using for in loop to populate the contents of the first row in the table:
# looping 3 times for the three columns of the first row of the table
# make it in for loop:
# range() => starts with 0
# 3 columns => 3 times to iterate
# range(3) => 0, 1, 2
for column in range(3): # => column=0, column=1, column=2 
   table3.rows[0].cells[column].text = t_headings[column] 

# save()


# looping the table rows:
# each row has first name, last name, program
for f_name, l_name, program in rows:
    cells = table3.add_row().cells
    cells[0].text =f_name
    cells[1].text =l_name
    cells[2].text =program


save()

